import os
import tkinter as tk
from tkinter import filedialog
import subprocess
import docx2txt
import docx
import keyword
import string
import glob
from PIL import Image, ImageTk
from enum import Enum

class Color(Enum):
    YELLOW = 5
 
def search_folder(event=None):
    folder_path = folder_path_var.get()
    keywords = keyword_entry.get().split()
    matches_listbox.delete(0, tk.END)
    
    found_matches = False  # flag to track if any matches were found
    
    for window, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith(('.docx')):
                text = docx2txt.process(os.path.join(window, file))
                for keyword in keywords:
                    if keyword.lower() in text.lower():
                        matches_listbox.insert(tk.END, os.path.join(window, file))
                        break
            if found_matches:
                    break  # break inner loop if match was found
        if found_matches:
            break  # break outer loop if match was found
    
    if not found_matches:
        matches_listbox.insert(tk.END, "No matches found.")
                    
    folder_label.config(text="Selected folder: " + folder_path)
       
def open_file(event):
    selection = matches_listbox.curselection()
    if len(selection) > 0:
        doc_path = matches_listbox.get(selection[0])
        if os.path.exists(doc_path):
            #subprocess.Popen(["start", "", doc_path], shell=True)
            subprocess.Popen(["start", "", doc_path], shell=True)
            for keyword in keyword_entry.get().split():
                highlight_keyword(keyword, doc_path)

def reset_data(event=None):
    folder_path_var.set('')
    keyword_entry.delete(0, tk.END)
    matches_listbox.delete(0, tk.END)
    folder_label.config(text='')
    
def extract_text_from_docx(doc_path):
    doc = docx.Document(doc_path)
    text = '\n'.join([p.text for p in doc.paragraphs])
    return text

def highlight_keyword(keyword, doc_path):
    doc = docx.Document(doc_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if keyword.lower() in run.text.lower():
                font = run.font
                font.highlight_color = Color.YELLOW.value
    doc.save(doc_path)

    
    
# create tkinter window
window = tk.Tk()
window.title("Ky Search")
window.geometry("845x590")
window.minsize(400, 300)

# center the window on the screen
screen_width = window.winfo_screenwidth()
screen_height = window.winfo_screenheight()
x = int((screen_width - window.winfo_reqwidth()) / 2)
y = int((screen_height - window.winfo_reqheight()) / 2)
window.geometry("+{}+{}".format(x, y))

# create select folder button and label
folder_label = tk.Label(window, text="Select a folder:")
folder_label.grid(row=0, column=0, pady=5)
folder_path_var = tk.StringVar() # variable to hold selected folder path
select_folder_button = tk.Button(window, text="Browse", command=lambda: folder_path_var.set(filedialog.askdirectory()))
select_folder_button.grid(row=1, column=0, pady=5)
reset_button = tk.Button(window, text="Reset", command=reset_data)
reset_button.grid(row=1, column=1, pady=5)
selected_folder_label = tk.Label(window, textvariable=folder_path_var)
selected_folder_label.grid(row=2, column=0, pady=5)

# create keyword entry widget and label
keyword_label = tk.Label(window, text="Enter keywords (separated by spaces):")
keyword_label.grid(row=3, column=0, pady=5)
keyword_entry = tk.Entry(window)
keyword_entry.grid(row=4, column=0, pady=5)
search_button = tk.Button(window, text="Search", command=search_folder)
search_button.grid(row=4, column=1, pady=5, padx=5)
execute_button = tk.Button(window, text="Close", command=window.destroy)
execute_button.grid(row=8, column=1, pady=5)

keyword_entry.bind('<Return>', search_folder)

# create matches listbox and label
matches_label = tk.Label(window, text="Matching files:")
matches_label.grid(row=5, column=0, pady=5)
matches_frame = tk.Frame(window, height=150)
matches_frame.grid(row=6, column=0, columnspan=2, pady=5, padx=5, sticky='nsew')
matches_frame.grid_rowconfigure(0, weight=1)
matches_frame.grid_columnconfigure(0, weight=1)
matches_listbox = tk.Listbox(matches_frame)
matches_listbox.pack(fill=tk.BOTH, expand=True)
matches_listbox.bind('<Double-Button-1>', open_file)

# center the window on the screen
screen_width = window.winfo_screenwidth()
screen_height = window.winfo_screenheight()
x = int((screen_width - window.winfo_reqwidth()) / 2)
y = int((screen_height - window.winfo_reqheight()) / 2)
window.geometry("+{}+{}".format(x, y))

# adjust the size of the listbox frame
window.grid_columnconfigure(0, weight=1)
window.grid_rowconfigure(6, weight=1)

# start the main event loop
window.mainloop()